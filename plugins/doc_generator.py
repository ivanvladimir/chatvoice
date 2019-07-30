#!/usr/bin/env python
# -*- coding: utf-8 -*-
#
# Ivan Vladimir Meza Ruiz 2018
# GPL 3.0

import random
import spacy
from docx import Document
from docx.shared import Inches

def execute(*args):
    nom=args[0]
    apP=args[1]
    apM=args[2]
    fec=args[3]
    dom=args[4]
    nac=args[5]
    civ=args[6]
    tel=args[7]
    ema=args[8]
    d_e=args[9]
    l_e=args[10]
    lug=args[11]
    idio=args[12]
    offi=args[13]
    f = open('plantilla.docx', 'rb')
    document = Document(f)
    f.close()
    p=document.add_paragraph('Datos personales', style='Intense Quote')
    n=document.add_paragraph('')
    n.add_run('Nombre: ').bold = True
    nn=nom+" "+apP+" "+apM
    nn=nn.title()
    n.add_run(nn)
    f=document.add_paragraph('')
    f.add_run('Fecha de nacimiento: ').bold = True
    f.add_run(fec)
    d=document.add_paragraph('')
    d.add_run('Domicilio: ').bold = True
    d.add_run(dom)
    na=document.add_paragraph('')
    na.add_run('Nancionalidad: ').bold = True
    na.add_run(nac)
    e=document.add_paragraph('')
    e.add_run('Estado civil: ').bold = True
    e.add_run(civ)
    t=document.add_paragraph('')
    t.add_run('Teléfono: ').bold = True
    t.add_run(tel)
    c=document.add_paragraph('')
    c.add_run('e-mail: ').bold = True
    c.add_run(ema)
    p2=document.add_paragraph('Formación académica', style='Intense Quote')
    de=document.add_paragraph('')
    de.add_run('Datos escolares: ').bold = True
    de.add_run(d_e)
    le=document.add_paragraph('')
    le.add_run('Proyectos escolares: ').bold = True
    table = document.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Nombre del proyecto'
    hdr_cells[1].text = 'Descripción'
    for nomp, desc in l_e:
        row_cells = table.add_row().cells
        row_cells[0].text = nomp
        row_cells[1].text = desc
    p3=document.add_paragraph('Experiencia laboral', style='Intense Quote')
    tablet = document.add_table(rows=1, cols=2)
    hdr_cellst = tablet.rows[0].cells
    hdr_cellst[0].text = 'Lugar y puesto'
    hdr_cellst[1].text = 'Periodo'
    for nompt, desct in lug:
        row_cellst = tablet.add_row().cells
        row_cellst[0].text = nompt
        row_cellst[1].text = desct
    p4=document.add_paragraph('Habilidades técnicas', style='Intense Quote')
    idi=document.add_paragraph('')
    idi.add_run('Idiomas: ').bold = True
    idi.add_run(idio)
    her=document.add_paragraph('')
    her.add_run('Herramientas Office: ').bold = True
    her.add_run(offi)
    document.save('curriculumHAUNTER.docx')

    return 'say "¡CURRICULUM GENERADO CON ÉXITO!"'
